import sys, io, shutil, json
from lxml import etree
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Namespaces ──────────────────────────────────────────────────────────────
M_NS      = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
_m        = lambda tag: f'{{{M_NS}}}{tag}'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

shutil.copy('SUSUN-SKRIPSI-ARIF_BACKUP.docx', 'SUSUN-SKRIPSI-ARIF.docx')
print("Restored from backup.")

doc = Document('SUSUN-SKRIPSI-ARIF.docx')
TNR = 'Arial'

# ═══════════════════════════════════════════════════════════════════════════════
# OMML HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def mt(text):
    """m:r > m:t plain text element."""
    r = etree.Element(_m('r'))
    t = etree.SubElement(r, _m('t'))
    t.set(XML_SPACE, 'preserve')
    t.text = text
    return r

def msub(base, sub):
    """Subscript: base_sub."""
    s = etree.Element(_m('sSub'))
    e = etree.SubElement(s, _m('e'));   e.append(mt(base))
    sb = etree.SubElement(s, _m('sub')); sb.append(mt(sub))
    return s

def mfrac(num_items, den_items):
    """Fraction num/den."""
    f = etree.Element(_m('f'))
    n = etree.SubElement(f, _m('num'))
    d = etree.SubElement(f, _m('den'))
    for item in (num_items if isinstance(num_items, list) else [num_items]):
        n.append(item)
    for item in (den_items if isinstance(den_items, list) else [den_items]):
        d.append(item)
    return f

def eq_ins(math_elements, after_el):
    """Insert centered OMML math paragraph after after_el. Returns new element."""
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'),'0'); sp.set(qn('w:after'),'0')
    sp.set(qn('w:line'),'360'); sp.set(qn('w:lineRule'),'auto')
    pPr.append(sp); new_p.append(pPr)

    mathPara = etree.Element(_m('oMathPara'))
    mpp = etree.SubElement(mathPara, _m('oMathParaPr'))
    jc_m = etree.SubElement(mpp, _m('jc'))
    jc_m.set(_m('val'), 'center')
    oMath = etree.SubElement(mathPara, _m('oMath'))
    for el in math_elements:
        oMath.append(el)
    new_p.append(mathPara)
    after_el.addnext(new_p)
    return new_p

# ═══════════════════════════════════════════════════════════════════════════════
# PARAGRAPH & TABLE HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
cur = None

def p(text, bold=False, italic=False, center=False, size=10, after=None, ind=-1):
    # ind=-1 = auto: bold/center → 0 (heading), else → 1800 (body text)
    # ind=0  = no indent (explicit)
    # ind>0  = raw twips
    actual_ind = (0 if (bold or center) else 1800) if ind == -1 else ind
    new_p = OxmlElement('w:p')
    pPr   = OxmlElement('w:pPr')
    s = OxmlElement('w:pStyle'); s.set(qn('w:val'), 'ListParagraph'); pPr.append(s)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center' if center else 'both'); pPr.append(jc)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'),'0'); sp.set(qn('w:after'),'0')
    sp.set(qn('w:line'),'360');  sp.set(qn('w:lineRule'),'auto'); pPr.append(sp)
    if actual_ind > 0:
        i2 = OxmlElement('w:ind'); i2.set(qn('w:left'), str(actual_ind)); pPr.append(i2)
    new_p.append(pPr)
    if text:
        r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),TNR); rf.set(qn('w:hAnsi'),TNR); rPr.append(rf)
        if bold:   rPr.append(OxmlElement('w:b'))
        if italic: rPr.append(OxmlElement('w:i'))
        sz = OxmlElement('w:sz');   sz.set(qn('w:val'),str(size*2));   rPr.append(sz)
        szc= OxmlElement('w:szCs'); szc.set(qn('w:val'),str(size*2));  rPr.append(szc)
        r.append(rPr)
        t = OxmlElement('w:t'); t.set(XML_SPACE,'preserve')
        t.text = text; r.append(t); new_p.append(r)
    ref = after if after is not None else cur
    ref.addnext(new_p)
    return new_p

def make_tbl(rows_data, hdr_rows=1):
    nc  = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=nc)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows_data):
        for ci, ci_data in enumerate(row):
            if isinstance(ci_data, dict):
                txt  = ci_data.get('t','')
                bold = ci_data.get('b', ri < hdr_rows)
                left = ci_data.get('l', False)
            else:
                txt  = str(ci_data)
                bold = ri < hdr_rows
                left = False
            cell = tbl.cell(ri, ci)
            cell.text = ''
            pp  = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER
            run = pp.add_run(txt)
            run.bold = bold; run.font.name = TNR; run.font.size = Pt(10)
            if ri < hdr_rows:
                tc   = cell._element
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
                shd.set(qn('w:fill'),'D9D9D9'); tcPr.append(shd)
    el = tbl._element; el.getparent().remove(el)
    return el

def ins_tbl(after_el, rows_data, caption, source='Hasil Pengolahan Data (2025)', hdr_rows=1):
    c = p(caption, bold=True, center=True, after=after_el)
    tbl_el = make_tbl(rows_data, hdr_rows)
    c.addnext(tbl_el)
    s = p('Sumber: ' + source, italic=True, center=True, after=tbl_el)
    e = p('', after=s)
    return e

def tbl_cond(rows, caption):
    HDR = ['Nilai','Jml MASUK','Jml TDK MSK','P(X|MASUK)','P(X|TDK MSK)','Hasil MASUK','Hasil TDK MSK']
    d = [HDR]
    for r in rows:
        d.append([{'t':r[0],'l':True}, r[1], r[2], r[3], r[4], r[5], r[6]])
    return d, caption

def fmt(val):
    v = str(val).replace('_',' ')
    if 'Rp.4.000.001' in v: return 'Rp4-6jt'
    if 'Rp.2.000.001' in v: return 'Rp2-4jt'
    if 'Diatas Rp.6' in v or '> Rp6' in v: return '>Rp6jt'
    if 'Rp.1.000.001' in v or ('Rp.1.000.000' in v and 'Rp.2.000.000' in v): return 'Rp1-2jt'
    if 'Dibawah' in v: return '<Rp1jt'
    return v

# ─── Load data ────────────────────────────────────────────────────────────────
with open('testing_results.json', encoding='utf-8') as f:
    jd = json.load(f)
results     = jd['results']
train_rows  = jd['train_rows']
n_train     = jd['n_train']
n_test      = jd['n_test']
prior_data  = jd['prior']
cond_tables = jd['cond_tables']
tp, tn, fp, fn = jd['tp'], jd['tn'], jd['fp'], jd['fn']
acc, prec, rec, f1 = jd['acc'], jd['prec'], jd['rec'], jd['f1']

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 1: INSERT IMAGES AT GMABAR PLACEHOLDERS
# ═══════════════════════════════════════════════════════════════════════════════
def replace_with_image(para_obj, img_path, caption_text):
    p_elem = para_obj._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); p_elem.insert(0, pPr)
    jc_el = pPr.find(qn('w:jc'))
    if jc_el is None:
        jc_el = OxmlElement('w:jc'); pPr.append(jc_el)
    jc_el.set(qn('w:val'), 'center')
    sp_el = pPr.find(qn('w:spacing'))
    if sp_el is None:
        sp_el = OxmlElement('w:spacing'); pPr.append(sp_el)
    sp_el.set(qn('w:before'),'0'); sp_el.set(qn('w:after'),'0')
    sp_el.set(qn('w:line'),'360'); sp_el.set(qn('w:lineRule'),'auto')

    run = para_obj.add_run()
    run.add_picture(img_path, width=Cm(14))

    # Caption paragraph
    cap_p = OxmlElement('w:p')
    cap_pPr = OxmlElement('w:pPr')
    cap_jc = OxmlElement('w:jc'); cap_jc.set(qn('w:val'), 'center'); cap_pPr.append(cap_jc)
    cap_sp = OxmlElement('w:spacing')
    cap_sp.set(qn('w:before'),'0'); cap_sp.set(qn('w:after'),'0')
    cap_sp.set(qn('w:line'),'360'); cap_sp.set(qn('w:lineRule'),'auto')
    cap_pPr.append(cap_sp); cap_p.append(cap_pPr)
    cap_r = OxmlElement('w:r')
    cap_rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'Arial'); rf.set(qn('w:hAnsi'),'Arial'); cap_rPr.append(rf)
    cap_rPr.append(OxmlElement('w:b'))
    cap_sz = OxmlElement('w:sz');   cap_sz.set(qn('w:val'),'20'); cap_rPr.append(cap_sz)
    cap_szc = OxmlElement('w:szCs'); cap_szc.set(qn('w:val'),'20'); cap_rPr.append(cap_szc)
    cap_r.append(cap_rPr)
    cap_t = OxmlElement('w:t'); cap_t.text = caption_text; cap_r.append(cap_t)
    cap_p.append(cap_r)
    p_elem.addnext(cap_p)
    return cap_p

img_count = 0
for para in doc.paragraphs:
    stripped = para.text.strip()
    if stripped == 'Gmabar' and img_count == 0:
        replace_with_image(para, 'Proses Bisnis Lama.png', 'Gambar 4.1 BPMN Proses Bisnis Lama')
        img_count += 1
        print("Image 1 inserted: Proses Bisnis Lama")
    elif stripped == 'Gambar' and img_count == 1:
        replace_with_image(para, 'Proses Bisnis Baru.png', 'Gambar 4.2 BPMN Proses Bisnis Baru')
        img_count += 1
        print("Image 2 inserted: Proses Bisnis Baru")
    if img_count >= 2:
        break
print(f"Total images inserted: {img_count}")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 2: FIND "HASIL ANALISIS METODE" PARAGRAPH (by text)
# ═══════════════════════════════════════════════════════════════════════════════
cur = None
for para in doc.paragraphs:
    if 'Hasil Anali' in para.text and 'Metode' in para.text:
        cur = para._element
        print(f"Found: '{para.text.strip()}'")
        break
if cur is None:
    raise RuntimeError("Tidak menemukan paragraf 'Hasil Analisis Metode'!")

# Delete 3 placeholder paragraphs right after it
deleted = 0
for _ in range(3):
    nxt = cur.getnext()
    if nxt is not None and nxt.tag == qn('w:p'):
        nxt.getparent().remove(nxt)
        deleted += 1
print(f"Deleted {deleted} placeholder paragraphs.")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 3: INSERT "HASIL ANALISIS METODE" CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

cur = p('Metode yang diterapkan dalam proses penelitian ini menggunakan algoritma Naive Bayes. '
        'Metode ini diharapkan dapat menghasilkan rekomendasi tindak lanjut konfirmasi pendaftaran '
        'mahasiswa baru secara efektif dan akurat. Berikut adalah beberapa langkah dengan pendekatan '
        'algoritma Naive Bayes:', after=cur)

# ── 1) Menyiapkan Dataset ────────────────────────────────────────────────────
cur = p('1)  Menyiapkan Dataset', bold=True, after=cur, ind=1800)
cur = p(f'Pada langkah ini dilakukan pengumpulan data historis pendaftaran mahasiswa baru Universitas Tazkia. '
        f'Data yang digunakan berjumlah 1.285 record yang selanjutnya dibagi menjadi dua dataset, '
        f'terdiri dari 80% untuk data training yaitu sebanyak {n_train} data, serta 20% untuk data testing '
        f'yaitu sebanyak {n_test} data. Berikut adalah sebagian dataset untuk data training dan testing yang dapat '
        f'dilihat pada Tabel 4.1 dan Tabel 4.2 di bawah ini.', after=cur)

HDR_TR = ['No','Nama','Kat. Jarak','Follow Up','Status Test','Nilai Test','Penghasilan','Kelas']
tbl_tr = [HDR_TR]
for r in train_rows[:30]:
    tbl_tr.append([
        str(r['no']),
        {'t': r['nama'][:22], 'l': True},
        fmt(r['vals'][0]), fmt(r['vals'][1]), fmt(r['vals'][2]),
        fmt(r['vals'][3]),
        {'t': fmt(r['vals'][4]), 'l': True},
        r['kelas'],
    ])
tbl_tr.append([{'t': '...', 'l': True}, {'t': '...', 'l': True},
               '...','...','...','...','...','...'])
tbl_tr.append([{'t': f'(dan seterusnya - total {n_train} data training)', 'b': True, 'l': True},
               '','','','','','',''])
cur = ins_tbl(cur, tbl_tr, 'Tabel 4.1 Dataset Untuk Data Training',
              'Data Internal Pendaftaran Universitas Tazkia')

HDR_TE = ['No','Nama','Kat. Jarak','Follow Up','Status Test','Nilai Test','Penghasilan','Prediksi']
tbl_te_rows = [HDR_TE]
for i, r in enumerate(results[:30], 1):
    tbl_te_rows.append([
        str(i), {'t': r['nama'][:22], 'l': True},
        fmt(r['vals'][0]), fmt(r['vals'][1]), fmt(r['vals'][2]),
        fmt(r['vals'][3]), {'t': fmt(r['vals'][4]), 'l': True}, '?'
    ])
tbl_te_rows.append([{'t': '...', 'l': True}, {'t': '...', 'l': True},
                    '...','...','...','...','...','...'])
tbl_te_rows.append([{'t': f'(dan seterusnya - total {n_test} data testing)', 'b': True, 'l': True},
                    '','','','','','',''])
cur = ins_tbl(cur, tbl_te_rows, 'Tabel 4.2 Dataset Untuk Data Testing',
              'Data Internal Pendaftaran Universitas Tazkia')

# ── 2) Pre-processing ─────────────────────────────────────────────────────────
cur = p('2)  Pre-processing Data', bold=True, after=cur, ind=1800)
cur = p('Pada langkah ini dilakukan pre-processing atau mempersiapkan data sebelum digunakan '
        'dalam algoritma Naive Bayes melalui pembersihan data dan seleksi fitur.', after=cur)
cur = p('a.  Pembersihan Data', bold=True, after=cur, ind=2520)
cur = p('Dilakukan pemeriksaan terhadap data yang memiliki nilai kosong (missing value) atau '
        'tidak relevan (noise). Berdasarkan hasil pemeriksaan, data pada Tabel 4.1 dan Tabel 4.2 '
        'tidak ditemukan nilai kosong sehingga seluruh data dapat langsung digunakan.', after=cur)
cur = p('b.  Seleksi Fitur', bold=True, after=cur, ind=2520)
cur = p('Seleksi fitur dilakukan untuk memilih atribut yang akan digunakan dalam proses '
        'klasifikasi. Hasil seleksi fitur dapat dilihat pada Tabel 4.3 berikut.', after=cur)

tbl_f = [
    ['Fitur/Atribut yang Ada Sebelumnya','Fitur/Atribut yang Digunakan'],
    [{'t':'Kategori Jarak Asal','l':True},      {'t':'Kategori Jarak Asal','l':True}],
    [{'t':'Kategori Asal Sekolah','l':True},    {'t':'- (tidak digunakan)','l':True}],
    [{'t':'Waktu Pendaftaran','l':True},         {'t':'- (tidak digunakan)','l':True}],
    [{'t':'Tingkat Follow Up Internal','l':True},{'t':'Tingkat Follow Up Internal','l':True}],
    [{'t':'Status Uang Pendaftaran','l':True},   {'t':'- (tidak digunakan)','l':True}],
    [{'t':'Status Test','l':True},               {'t':'Status Test','l':True}],
    [{'t':'Kategori Nilai Test','l':True},       {'t':'Kategori Nilai Test','l':True}],
    [{'t':'Kategori Penghasilan','l':True},      {'t':'Kategori Penghasilan','l':True}],
    [{'t':'Status Retensi Final (Target)','l':True},{'t':'Status Retensi Final (Target)','l':True}],
]
cur = ins_tbl(cur, tbl_f, 'Tabel 4.3 Hasil Seleksi Fitur')

# ── 3) Variabel ───────────────────────────────────────────────────────────────
cur = p('3)  Menentukan Variabel Penelitian', bold=True, after=cur, ind=1800)
cur = p('Berdasarkan hasil seleksi fitur, variabel yang digunakan dalam proses klasifikasi '
        'Naive Bayes dapat dilihat pada Tabel 4.4 berikut.', after=cur)

tbl_v = [
    ['Kode','Atribut','Tipe','Nilai'],
    ['X1',{'t':'Kategori Jarak Asal','l':True},      {'t':'Variabel Atribut','l':True},{'t':'Dekat, Sedang, Jauh','l':True}],
    ['X2',{'t':'Tingkat Follow Up Internal','l':True},{'t':'Variabel Atribut','l':True},{'t':'Belum Dihubungi, Kontak Awal, Follow Up, Follow Up Intensif','l':True}],
    ['X3',{'t':'Status Test','l':True},               {'t':'Variabel Atribut','l':True},{'t':'Sudah Tes, Belum Tes','l':True}],
    ['X4',{'t':'Kategori Nilai Test','l':True},       {'t':'Variabel Atribut','l':True},{'t':'Nilai Tinggi, Nilai Sedang, Tidak Ada Nilai','l':True}],
    ['X5',{'t':'Kategori Penghasilan','l':True},      {'t':'Variabel Atribut','l':True},{'t':'>Rp6jt, Rp4-6jt, Rp2-4jt, Rp1-2jt, <Rp1jt, Tidak Diketahui','l':True}],
    ['Y', {'t':'Status Retensi Final','l':True},      {'t':'Kelas Target','l':True},    {'t':'MASUK, TIDAK MASUK','l':True}],
]
cur = ins_tbl(cur, tbl_v, 'Tabel 4.4 Variabel Penelitian')

# ── 4) Prior Probability ──────────────────────────────────────────────────────
cur = p('4)  Menghitung Prior Probability', bold=True, after=cur, ind=1800)
cur = p(f'Pada tahap ini dilakukan perhitungan probabilitas prior atau peluang kemunculan setiap '
        f'kelas target (Y) yaitu MASUK dan TIDAK MASUK, dihitung berdasarkan data training pada '
        f'Tabel 4.1 yang berjumlah {n_train} data. Rumus prior probability adalah sebagai berikut:', after=cur)

# Rumus Prior: P(Yi) = ni/n
cur = eq_ins([
    mt('P('), msub('Y','i'), mt(') = '),
    mfrac([msub('n','i')], [mt('n')])
], after_el=cur)

cur = p('Sehingga diperoleh nilai prior probability untuk masing-masing kelas sebagai berikut:', after=cur)

n_m = prior_data['MASUK']['n']
n_t = prior_data['TIDAK MASUK']['n']
p_m = prior_data['MASUK']['prob']
p_t = prior_data['TIDAK MASUK']['prob']

cur = eq_ins([
    mt(f'P(Y = MASUK) = '),
    mfrac([mt(str(n_m))], [mt(str(n_train))]),
    mt(f' = {str(p_m).replace(".", ",")}')
], after_el=cur)

cur = eq_ins([
    mt(f'P(Y = TIDAK MASUK) = '),
    mfrac([mt(str(n_t))], [mt(str(n_train))]),
    mt(f' = {str(p_t).replace(".", ",")}')
], after_el=cur)

cur = p('Berdasarkan hasil perhitungan di atas, maka hasil perhitungan prior probability dapat '
        'dilihat pada Tabel 4.5.', after=cur)

tbl_pr = [
    ['Kelas (Y)','Jumlah Kemunculan (ni)','P(Y)','Hasil'],
    [{'t':'MASUK','l':True},       str(n_m), f'{n_m} / {n_train}', str(p_m).replace('.',',')],
    [{'t':'TIDAK MASUK','l':True}, str(n_t), f'{n_t} / {n_train}', str(p_t).replace('.',',')],
    [{'t':f'Total Data (n)','b':True,'l':True}, str(n_train),'',''],
]
cur = ins_tbl(cur, tbl_pr, 'Tabel 4.5 Hasil Perhitungan Prior Probability')

# ── 5) Likelihood ─────────────────────────────────────────────────────────────
cur = p('5)  Menghitung Likelihood (Conditional Probability)', bold=True, after=cur, ind=1800)
cur = p('Pada tahap ini dilakukan perhitungan likelihood atau conditional probability, yaitu '
        'probabilitas kemunculan setiap nilai atribut pada masing-masing kelas berdasarkan '
        'data training. Rumus yang digunakan dengan Laplace Smoothing:', after=cur)

# Rumus Laplace: P(Xi|H) = (ni+1)/(n+K)
cur = eq_ins([
    mt('P('), msub('X','i'), mt(' | H) = '),
    mfrac(
        [msub('n','i'), mt(' + 1')],
        [mt('n + K')]
    )
], after_el=cur)

cur = p('Keterangan: ni = jumlah data bernilai Xi pada kelas H; n = total data kelas H; '
        'K = jumlah nilai unik atribut tersebut. '
        'Penambahan nilai 1 pada pembilang dan K pada penyebut merupakan teknik Laplace Smoothing '
        'untuk menghindari nilai probabilitas nol.', after=cur)

def build_cond_rows(fname):
    ft = cond_tables[fname]
    rows = []
    total_m, total_t = 0, 0
    for v in ft['vals']:
        m = ft['MASUK'][v]
        t = ft['TIDAK MASUK'][v]
        total_m += m['ni']; total_t += t['ni']
        rows.append((fmt(v), str(m['ni']), str(t['ni']),
                     m['frac'], t['frac'],
                     str(m['prob']).replace('.',','),
                     str(t['prob']).replace('.',',')))
    rows.append(('Total', str(total_m), str(total_t), '', '', '', ''))
    return rows

d, c = tbl_cond(build_cond_rows('kategori_jarak_asal'),
                'Tabel 4.6 Hasil Perhitungan Conditional Probability Kategori Jarak Asal')
cur = ins_tbl(cur, d, c)

d, c = tbl_cond(build_cond_rows('tingkat_follow_up_internal'),
                'Tabel 4.7 Hasil Perhitungan Conditional Probability Tingkat Follow Up Internal')
cur = ins_tbl(cur, d, c)

d, c = tbl_cond(build_cond_rows('status_test'),
                'Tabel 4.8 Hasil Perhitungan Conditional Probability Status Test')
cur = ins_tbl(cur, d, c)

d, c = tbl_cond(build_cond_rows('kategori_nilai_test'),
                'Tabel 4.9 Hasil Perhitungan Conditional Probability Kategori Nilai Test')
cur = ins_tbl(cur, d, c)

d, c = tbl_cond(build_cond_rows('kategori_penghasilan'),
                'Tabel 4.10 Hasil Perhitungan Conditional Probability Kategori Penghasilan')
cur = ins_tbl(cur, d, c)

# ── 6) Posterior Probability ──────────────────────────────────────────────────
cur = p('6)  Menghitung Posterior Probability', bold=True, after=cur, ind=1800)
cur = p('Pada tahapan ini dilakukan perhitungan posterior probability untuk memprediksi kelas Y '
        '(MASUK dan TIDAK MASUK) berdasarkan seluruh data testing pada Tabel 4.2. '
        'Rumus yang digunakan:', after=cur)

# Rumus Posterior: P(Yi|X) = P(Xi|Yi) x P(Yi)
cur = eq_ins([
    mt('P('), msub('Y','i'), mt('|X) = P('), msub('X','i'), mt('|'), msub('Y','i'), mt(') × P('), msub('Y','i'), mt(')')
], after_el=cur)

cur = p(f'Nilai dari prior probability dan conditional probability yang digunakan pada perhitungan '
        f'ini diambil dari Tabel 4.5 sampai dengan Tabel 4.10. Hasil perhitungan posterior '
        f'probability untuk {n_test} data testing dapat dilihat pada Tabel 4.11 berikut '
        f'(ditampilkan 15 data sebagai sampel).', after=cur)

tbl_post = [['No','Nama','P(Y|X)','Kat. Jarak','Follow Up','Status Test','Nilai Test','Penghasilan','Prior','Posterior']]
for i, r in enumerate(results[:15], 1):
    pm_vals = [str(v) for v in r['pm']]
    pt_vals = [str(v) for v in r['pt']]
    tbl_post.append([
        str(i), {'t': r['nama'][:22], 'l': True}, 'P(MASUK|X)',
        pm_vals[0], pm_vals[1], pm_vals[2], pm_vals[3], pm_vals[4],
        str(r['prior_m']), f"{r['post_m']:.8f}"
    ])
    tbl_post.append([
        '', '', 'P(TDK MSK|X)',
        pt_vals[0], pt_vals[1], pt_vals[2], pt_vals[3], pt_vals[4],
        str(r['prior_t']), f"{r['post_t']:.8f}"
    ])
tbl_post.append([{'t': f'... (dan seterusnya, total {n_test} data testing)', 'b': True, 'l': True},
                 '','','','','','','','',''])
cur = ins_tbl(cur, tbl_post, 'Tabel 4.11 Hasil Perhitungan Posterior Probability')

# ── 7) Klasifikasi ────────────────────────────────────────────────────────────
cur = p('7)  Melakukan Klasifikasi (Prediksi)', bold=True, after=cur, ind=1800)
cur = p('Berdasarkan hasil perhitungan posterior probability pada Tabel 4.11, selanjutnya '
        'masing-masing nilai posterior dibandingkan untuk menentukan kelas prediksi, '
        'dengan ketentuan sebagai berikut:', after=cur)
cur = p('a.  P(MASUK|X)  >  P(TIDAK MASUK|X)  →  Prediksi = MASUK', after=cur, ind=2520)
cur = p('b.  P(MASUK|X)  <  P(TIDAK MASUK|X)  →  Prediksi = TIDAK MASUK', after=cur, ind=2520)
cur = p(f'Adapun hasil prediksi untuk seluruh {n_test} data testing dapat dilihat pada Tabel 4.12 berikut '
        f'(ditampilkan 30 data sebagai sampel).', after=cur)

tbl_pred = [['No','Nama','Nilai Posterior Terbesar','Status Aktual','Prediksi']]
for i, r in enumerate(results[:30], 1):
    tbl_pred.append([
        str(i), {'t': r['nama'][:22], 'l': True},
        f"{max(r['post_m'], r['post_t']):.8f}",
        r['aktual'], r['pred']
    ])
tbl_pred.append([{'t': f'... (dan seterusnya, total {n_test} data testing)', 'b': True, 'l': True},
                 '','','',''])
cur = ins_tbl(cur, tbl_pred, 'Tabel 4.12 Hasil Prediksi Status Konfirmasi Pendaftaran')

# ── 8) Uji Hasil ──────────────────────────────────────────────────────────────
cur = p('8)  Uji Hasil', bold=True, after=cur, ind=1800)
cur = p(f'Tahapan ini dilakukan untuk mengevaluasi kinerja model menggunakan Confusion Matrix. '
        f'Evaluasi dilakukan dengan membandingkan hasil prediksi algoritma Naive Bayes terhadap '
        f'status aktual pada seluruh {n_test} data testing. '
        f'Hasil evaluasi model secara lengkap dibahas pada bagian C. Pembahasan.', after=cur)

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 4: INSERT d. ANALISIS KEBUTUHAN SISTEM (Use Case Diagram)
# ═══════════════════════════════════════════════════════════════════════════════
aks_cur = None
for para in doc.paragraphs:
    if para.text.strip() == 'Analisis Kebutuhan Sistem':
        aks_cur = para._element
        print("Found: 'Analisis Kebutuhan Sistem'")
        break

if aks_cur is not None:
    # Intro
    aks_cur = p('Pemodelan objek pada aplikasi yang akan dikembangkan dijelaskan dalam bentuk '
                'diagram use case berdasarkan pada proses rekomendasi tindak lanjut konfirmasi '
                'pendaftaran mahasiswa baru, untuk memodelkan serta mengorganisasi pada aplikasi '
                'sehingga mendapatkan keluaran aplikasi sesuai dengan yang diharapkan dan '
                'dibutuhkan oleh pengguna. Berikut diagram use case pada aplikasi yang akan '
                'dikembangkan dapat dilihat pada Gambar 4.3.', after=aks_cur)

    # Gambar Use Case Diagram
    img_p_elem = OxmlElement('w:p')
    img_pPr = OxmlElement('w:pPr')
    img_jc = OxmlElement('w:jc'); img_jc.set(qn('w:val'), 'center'); img_pPr.append(img_jc)
    img_sp = OxmlElement('w:spacing')
    img_sp.set(qn('w:before'),'0'); img_sp.set(qn('w:after'),'0')
    img_sp.set(qn('w:line'),'360'); img_sp.set(qn('w:lineRule'),'auto')
    img_pPr.append(img_sp); img_p_elem.append(img_pPr)
    aks_cur.addnext(img_p_elem)

    from docx.oxml import OxmlElement as OE
    from docx.shared import Cm as DCm
    # Find the paragraph object to add picture
    for para in doc.paragraphs:
        if para._element is img_p_elem:
            run = para.add_run()
            run.add_picture('Use Case Diagram.png', width=DCm(15))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    aks_cur = img_p_elem

    # Caption
    cap_p = OxmlElement('w:p')
    cap_pPr = OxmlElement('w:pPr')
    cap_jc = OxmlElement('w:jc'); cap_jc.set(qn('w:val'), 'center'); cap_pPr.append(cap_jc)
    cap_sp = OxmlElement('w:spacing')
    cap_sp.set(qn('w:before'),'0'); cap_sp.set(qn('w:after'),'0')
    cap_sp.set(qn('w:line'),'360'); cap_sp.set(qn('w:lineRule'),'auto')
    cap_pPr.append(cap_sp); cap_p.append(cap_pPr)
    cap_r = OxmlElement('w:r')
    cap_rPr = OxmlElement('w:rPr')
    cap_rf = OxmlElement('w:rFonts'); cap_rf.set(qn('w:ascii'),'Arial'); cap_rf.set(qn('w:hAnsi'),'Arial'); cap_rPr.append(cap_rf)
    cap_rPr.append(OxmlElement('w:b'))
    cap_sz = OxmlElement('w:sz'); cap_sz.set(qn('w:val'),'20'); cap_rPr.append(cap_sz)
    cap_szc = OxmlElement('w:szCs'); cap_szc.set(qn('w:val'),'20'); cap_rPr.append(cap_szc)
    cap_r.append(cap_rPr)
    cap_t = OxmlElement('w:t'); cap_t.text = 'Gambar 4.3 Diagram Use Case'; cap_r.append(cap_t)
    cap_p.append(cap_r)
    aks_cur.addnext(cap_p)
    aks_cur = cap_p

    aks_cur = p('Sumber: Hasil Perancangan (2025)', italic=True, center=True, after=aks_cur)

    # Penjelasan
    aks_cur = p('Pada Gambar 4.3 dijelaskan bahwa terdapat 1 (satu) aktor dalam aplikasi '
                'rekomendasi tindak lanjut konfirmasi pendaftaran mahasiswa baru yaitu Tim Marketing. '
                'Tim Marketing diharuskan untuk login terlebih dahulu agar dapat mengakses aplikasi '
                'tersebut. Adapun use case yang dapat dilakukan oleh Tim Marketing adalah sebagai '
                'berikut:', after=aks_cur)

    aks_cur = p('a.  Login: Tim Marketing melakukan autentikasi untuk masuk ke dalam sistem.',
                after=aks_cur, ind=2520)
    aks_cur = p('b.  Logout: Tim Marketing keluar dari sistem setelah selesai menggunakan aplikasi.',
                after=aks_cur, ind=2520)
    aks_cur = p('c.  Kelola Data Pendaftar: Tim Marketing dapat mengelola data calon mahasiswa '
                'baru yang terdaftar dalam sistem, termasuk melihat dan memperbarui data.',
                after=aks_cur, ind=2520)
    aks_cur = p('d.  Lihat Hasil Klasifikasi Naive Bayes: Tim Marketing dapat melihat hasil '
                'klasifikasi algoritma Naive Bayes terhadap seluruh data calon mahasiswa.',
                after=aks_cur, ind=2520)
    aks_cur = p('e.  Lihat Rekomendasi Tindak Lanjut: Tim Marketing dapat melihat rekomendasi '
                'tindak lanjut yang dihasilkan sistem untuk setiap calon mahasiswa berdasarkan '
                'hasil prediksi Naive Bayes.',
                after=aks_cur, ind=2520)
    aks_cur = p('f.  Lihat Detail Hasil Prediksi: Tim Marketing dapat melihat detail '
                'perhitungan probabilitas prediksi untuk setiap data calon mahasiswa.',
                after=aks_cur, ind=2520)
    aks_cur = p('g.  Cetak / Export Laporan: Tim Marketing dapat mencetak atau mengekspor '
                'laporan hasil rekomendasi tindak lanjut konfirmasi pendaftaran.',
                after=aks_cur, ind=2520)
    print("d. Analisis Kebutuhan Sistem inserted.")
else:
    print("WARNING: 'Analisis Kebutuhan Sistem' paragraph not found.")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5: DESIGN PRODUK — Activity Diagrams
# ═══════════════════════════════════════════════════════════════════════════════
def ins_img_para(after_el, img_path, caption_text, width_cm=13, source='Sumber: Hasil Perancangan (2025)'):
    """Insert centered image paragraph + caption + source after after_el. Returns last element."""
    # Image paragraph
    img_p = OxmlElement('w:p')
    img_pPr = OxmlElement('w:pPr')
    img_jc = OxmlElement('w:jc'); img_jc.set(qn('w:val'), 'center'); img_pPr.append(img_jc)
    img_sp = OxmlElement('w:spacing')
    img_sp.set(qn('w:before'),'0'); img_sp.set(qn('w:after'),'0')
    img_sp.set(qn('w:line'),'360'); img_sp.set(qn('w:lineRule'),'auto')
    img_pPr.append(img_sp); img_p.append(img_pPr)
    after_el.addnext(img_p)

    # Add picture via temporary paragraph
    tmp = doc.add_paragraph()
    run = tmp.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    for r in list(tmp._element.findall(qn('w:r'))):
        img_p.append(r)
    tmp._element.getparent().remove(tmp._element)

    # Caption
    cap = p(caption_text, bold=True, center=True, after=img_p)
    # Source
    src = p(source, italic=True, center=True, after=cap)
    return src

dp_cur = None
for para in doc.paragraphs:
    if para.text.strip() == 'Design Produk':
        dp_cur = para._element
        break

if dp_cur is not None:
    dp_cur = p('Berikut ini desain produk yang dikembangkan dalam sistem rekomendasi '
               'tindak lanjut konfirmasi pendaftaran mahasiswa baru menggunakan algoritma '
               'Naive Bayes, yaitu sebagai berikut:', after=dp_cur)

    dp_cur = p('a.  Activity Diagram', bold=True, after=dp_cur, ind=1800)
    dp_cur = p('Pada activity diagram ini dijelaskan alur aktivitas yang dilakukan oleh '
               'Tim Marketing dan alur yang dihasilkan oleh sistem, mulai dari proses awal '
               'membuka aplikasi hingga mendapatkan hasil rekomendasi tindak lanjut konfirmasi '
               'pendaftaran mahasiswa baru. Berikut adalah activity diagram yang dibuat dalam '
               'sistem rekomendasi tindak lanjut konfirmasi pendaftaran mahasiswa baru:', after=dp_cur)

    act_diagrams = [
        ('Activity Diagram Login.png',
         '1)  Activity Diagram Login',
         'Gambar 4.4 Activity Diagram Login',
         'Pada Gambar 4.4 terdapat activity diagram login yang diawali dengan Tim Marketing '
         'membuka aplikasi, kemudian sistem menampilkan halaman login. Selanjutnya Tim Marketing '
         'memasukkan username dan password lalu mengklik tombol login. Sistem memvalidasi '
         'kredensial tersebut. Apabila valid maka sistem menampilkan halaman utama, apabila '
         'tidak valid maka sistem menampilkan pesan error dan Tim Marketing diminta '
         'memasukkan kembali username dan password.'),

        ('Activity Diagram Logout.png',
         '2)  Activity Diagram Logout',
         'Gambar 4.5 Activity Diagram Logout',
         'Pada Gambar 4.5 terdapat activity diagram logout yang diawali dengan Tim Marketing '
         'mengklik tombol logout. Sistem kemudian menghapus sesi pengguna dan menampilkan '
         'kembali halaman login. Selanjutnya Tim Marketing diarahkan ke halaman login sebagai '
         'tanda bahwa proses logout telah berhasil dilakukan.'),

        ('Activity Diagram Kelola Data Pendaftar.png',
         '3)  Activity Diagram Kelola Data Pendaftar',
         'Gambar 4.6 Activity Diagram Kelola Data Pendaftar',
         'Pada Gambar 4.6 terdapat activity diagram kelola data pendaftar yang diawali dengan '
         'Tim Marketing membuka halaman data pendaftar dan sistem menampilkan daftar data yang '
         'ada. Tim Marketing mengklik tambah data, sistem menampilkan form, lalu Tim Marketing '
         'mengisi dan menyimpan data. Sistem memvalidasi data tersebut; jika valid maka data '
         'disimpan dan notifikasi ditampilkan, jika tidak valid maka sistem menampilkan pesan '
         'error dan Tim Marketing diminta mengisi ulang form.'),

        ('Activity Diagram Lihat Hasil Klasifikasi.png',
         '4)  Activity Diagram Lihat Hasil Klasifikasi Naive Bayes',
         'Gambar 4.7 Activity Diagram Lihat Hasil Klasifikasi Naive Bayes',
         'Pada Gambar 4.7 terdapat activity diagram lihat hasil klasifikasi Naive Bayes yang '
         'diawali dengan Tim Marketing membuka menu hasil klasifikasi. Sistem kemudian '
         'memproses klasifikasi menggunakan algoritma Naive Bayes dan menampilkan hasilnya. '
         'Selanjutnya Tim Marketing dapat melihat hasil klasifikasi yang telah diproses.'),

        ('Activity Diagram Lihat Rekomendasi.png',
         '5)  Activity Diagram Lihat Rekomendasi Tindak Lanjut',
         'Gambar 4.8 Activity Diagram Lihat Rekomendasi Tindak Lanjut',
         'Pada Gambar 4.8 terdapat activity diagram lihat rekomendasi tindak lanjut yang '
         'diawali dengan Tim Marketing membuka menu rekomendasi. Sistem melakukan proses '
         'generate rekomendasi tindak lanjut berdasarkan hasil klasifikasi Naive Bayes '
         'dan menampilkan hasilnya. Selanjutnya Tim Marketing dapat melihat rekomendasi '
         'tindak lanjut untuk setiap calon mahasiswa.'),

        ('Activity Diagram Lihat Detail Prediksi.png',
         '6)  Activity Diagram Lihat Detail Hasil Prediksi',
         'Gambar 4.9 Activity Diagram Lihat Detail Hasil Prediksi',
         'Pada Gambar 4.9 terdapat activity diagram lihat detail hasil prediksi yang '
         'diawali dengan Tim Marketing memilih data pendaftar yang ingin dilihat detailnya. '
         'Sistem mengambil detail hasil prediksi dan menampilkannya. Selanjutnya Tim Marketing '
         'dapat melihat detail perhitungan probabilitas Naive Bayes untuk calon mahasiswa '
         'yang dipilih.'),

        ('Activity Diagram Cetak Export Laporan.png',
         '7)  Activity Diagram Cetak / Export Laporan',
         'Gambar 4.10 Activity Diagram Cetak / Export Laporan',
         'Pada Gambar 4.10 terdapat activity diagram cetak atau export laporan yang diawali '
         'dengan Tim Marketing membuka menu laporan dan sistem menampilkan halaman laporan. '
         'Tim Marketing memilih rentang data lalu mengklik cetak atau export laporan. '
         'Sistem melakukan generate file laporan dalam format PDF atau Excel dan '
         'menyediakan file tersebut untuk diunduh oleh Tim Marketing.'),
    ]

    for (img_file, num_heading, caption, explanation) in act_diagrams:
        dp_cur = p(num_heading, bold=True, after=dp_cur, ind=2520)
        dp_cur = ins_img_para(dp_cur, img_file, caption, width_cm=15)
        dp_cur = p(explanation, after=dp_cur)

    print("Design Produk - Activity Diagrams inserted.")

    # ── b. Sequence Diagram ───────────────────────────────────────────────────
    dp_cur = p('b.  Sequence Diagram', bold=True, after=dp_cur, ind=1800)
    dp_cur = p('Pada sequence diagram ini dijelaskan alur interaksi antara Tim Marketing '
               'dengan komponen sistem secara berurutan sesuai garis waktunya, meliputi '
               'Form/Halaman antarmuka, Controller, dan Database. Berikut adalah sequence '
               'diagram yang dibuat dalam sistem rekomendasi tindak lanjut konfirmasi '
               'pendaftaran mahasiswa baru:', after=dp_cur)

    seq_diagrams = [
        ('Sequence Diagram Login.png',
         '1)  Sequence Diagram Login',
         'Gambar 4.11 Sequence Diagram Login',
         'Pada Gambar 4.11 terdapat sequence diagram login yang menggambarkan interaksi '
         'Tim Marketing dengan Form Login, Controller Login, Database, dan Halaman Utama. '
         'Tim Marketing membuka form login dan memasukkan username serta password. Form Login '
         'meneruskan proses ke Controller Login yang melakukan pengecekan kredensial ke Database. '
         'Apabila kredensial tidak valid, sistem menampilkan notifikasi gagal. Apabila valid, '
         'sistem menampilkan notifikasi berhasil dan mengarahkan Tim Marketing ke Halaman Utama.'),

        ('Sequence Diagram Logout.png',
         '2)  Sequence Diagram Logout',
         'Gambar 4.12 Sequence Diagram Logout',
         'Pada Gambar 4.12 terdapat sequence diagram logout yang menggambarkan interaksi '
         'Tim Marketing saat melakukan logout dari sistem. Tim Marketing mengklik tombol '
         'logout di Halaman Utama, kemudian Controller Logout memproses penghapusan sesi '
         'dan dikonfirmasi oleh Database. Setelah sesi berhasil dihapus, sistem '
         'mengarahkan Tim Marketing kembali ke Form Login.'),

        ('Sequence Diagram Kelola Data Pendaftar.png',
         '3)  Sequence Diagram Kelola Data Pendaftar',
         'Gambar 4.13 Sequence Diagram Kelola Data Pendaftar',
         'Pada Gambar 4.13 terdapat sequence diagram kelola data pendaftar yang menggambarkan '
         'interaksi Tim Marketing dalam mengelola data calon mahasiswa. Tim Marketing membuka '
         'halaman data pendaftar dan sistem mengambil data dari Database untuk ditampilkan. '
         'Tim Marketing dapat melakukan tambah, edit, atau hapus data yang kemudian diproses '
         'oleh Controller dan disimpan ke Database. Sistem menampilkan notifikasi sukses '
         'setelah operasi berhasil dilakukan.'),

        ('Sequence Diagram Lihat Hasil Klasifikasi.png',
         '4)  Sequence Diagram Lihat Hasil Klasifikasi',
         'Gambar 4.14 Sequence Diagram Lihat Hasil Klasifikasi',
         'Pada Gambar 4.14 terdapat sequence diagram lihat hasil klasifikasi yang menggambarkan '
         'proses melihat hasil klasifikasi Naive Bayes. Tim Marketing memilih data pendaftar '
         'dan sistem mengambil data tersebut dari Database. Controller kemudian menjalankan '
         'proses klasifikasi Naive Bayes secara internal, menyimpan hasilnya ke Database, '
         'dan menampilkan hasil klasifikasi kepada Tim Marketing.'),

        ('Sequence Diagram Lihat Rekomendasi.png',
         '5)  Sequence Diagram Lihat Rekomendasi',
         'Gambar 4.15 Sequence Diagram Lihat Rekomendasi',
         'Pada Gambar 4.15 terdapat sequence diagram lihat rekomendasi yang menggambarkan '
         'proses melihat rekomendasi tindak lanjut. Tim Marketing membuka halaman rekomendasi '
         'dan sistem mengambil data serta hasil klasifikasi dari Database. Controller '
         'melakukan proses generate rekomendasi secara internal dan menampilkan hasil '
         'rekomendasi tindak lanjut kepada Tim Marketing.'),

        ('Sequence Diagram Lihat Detail Prediksi.png',
         '6)  Sequence Diagram Lihat Detail Prediksi',
         'Gambar 4.16 Sequence Diagram Lihat Detail Prediksi',
         'Pada Gambar 4.16 terdapat sequence diagram lihat detail prediksi yang menggambarkan '
         'proses melihat detail prediksi untuk setiap pendaftar. Tim Marketing memilih data '
         'pendaftar yang ingin dilihat detailnya. Sistem mengambil detail data dan nilai '
         'probabilitas dari Database, kemudian menampilkan informasi detail prediksi '
         'Naive Bayes kepada Tim Marketing.'),

        ('Sequence Diagram Cetak Export Laporan.png',
         '7)  Sequence Diagram Cetak / Export Laporan',
         'Gambar 4.17 Sequence Diagram Cetak / Export Laporan',
         'Pada Gambar 4.17 terdapat sequence diagram cetak atau export laporan yang menggambarkan '
         'proses cetak atau export laporan. Tim Marketing memilih format laporan yang diinginkan '
         '(PDF atau Excel). Controller mengambil data dari Database, melakukan proses generate '
         'dokumen laporan secara internal, dan menyediakan file yang dapat diunduh atau '
         'dicetak oleh Tim Marketing.'),
    ]

    for (img_file, num_heading, caption, explanation) in seq_diagrams:
        dp_cur = p(num_heading, bold=True, after=dp_cur, ind=2520)
        dp_cur = ins_img_para(dp_cur, img_file, caption, width_cm=15)
        dp_cur = p(explanation, after=dp_cur)

    print("Design Produk - Sequence Diagrams inserted.")
else:
    print("WARNING: 'Design Produk' paragraph not found.")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5b: PERANCANGAN DESAIN ANTARMUKA (wireframe mockups)
# ═══════════════════════════════════════════════════════════════════════════════
pda_cur = None
for para in doc.paragraphs:
    if para.text.strip() == 'Perancang Design Antar Muka':
        if para.runs:
            para.runs[0].text = 'Perancangan Desain Antarmuka'
            for r in para.runs[1:]:
                r.text = ''
        else:
            para.add_run('Perancangan Desain Antarmuka')
        pda_cur = para._element
        break

if pda_cur is not None:
    pda_cur = p('Berdasarkan hasil analisis kebutuhan sistem dan design produk yang telah '
                'dijelaskan sebelumnya, tahap selanjutnya adalah merancang desain antarmuka '
                '(interface) aplikasi sebagai gambaran awal tampilan sistem sebelum masuk ke '
                'tahap pengkodean. Rancangan antarmuka ini dibuat agar Tim Marketing sebagai '
                'pengguna dapat memahami alur dan tata letak fitur pada sistem rekomendasi '
                'tindak lanjut konfirmasi pendaftaran mahasiswa baru. Berikut adalah rancangan '
                'desain antarmuka yang dibuat:', after=pda_cur)

    ui_pages = [
        ('Interface Login.png',
         '1)  Rancangan Halaman Login',
         'Gambar 4.18 Rancangan Antarmuka Halaman Login',
         'Gambar 4.18 menunjukkan rancangan halaman login yang menjadi titik masuk (entry point) '
         'aplikasi. Tim Marketing diharuskan memasukkan username dan password yang valid sebelum '
         'dapat mengakses fitur-fitur pada sistem.'),

        ('Interface Menu Utama.png',
         '2)  Rancangan Halaman Dashboard',
         'Gambar 4.19 Rancangan Antarmuka Halaman Dashboard',
         'Gambar 4.19 menunjukkan rancangan halaman dashboard yang tampil setelah Tim Marketing '
         'berhasil login. Halaman ini menampilkan ringkasan jumlah pendaftar, hasil prediksi '
         'MASUK/TIDAK MASUK, akurasi model, serta daftar prediksi terbaru.'),

        ('Interface Data Pendaftar.png',
         '3)  Rancangan Halaman Data Pendaftar',
         'Gambar 4.20 Rancangan Antarmuka Halaman Data Pendaftar',
         'Gambar 4.20 menunjukkan rancangan halaman data pendaftar yang menampilkan daftar seluruh '
         'calon mahasiswa beserta atribut yang digunakan dalam perhitungan Naive Bayes. Pada halaman '
         'ini Tim Marketing dapat mencari, menambah, mengubah, dan menghapus data pendaftar.'),

        ('Interface Tambah Edit Data Pendaftar.png',
         '4)  Rancangan Halaman Tambah/Ubah Data Pendaftar',
         'Gambar 4.21 Rancangan Antarmuka Halaman Tambah/Ubah Data Pendaftar',
         'Gambar 4.21 menunjukkan rancangan form input yang digunakan untuk menambah atau mengubah '
         'data pendaftar, meliputi kategori jarak asal, tingkat follow up internal, status test, '
         'kategori nilai test, dan kategori penghasilan.'),

        ('Interface Hasil Klasifikasi.png',
         '5)  Rancangan Halaman Hasil Klasifikasi Naive Bayes',
         'Gambar 4.22 Rancangan Antarmuka Halaman Hasil Klasifikasi Naive Bayes',
         'Gambar 4.22 menunjukkan rancangan halaman hasil klasifikasi yang menampilkan nilai '
         'probabilitas P(MASUK|X) dan P(TIDAK MASUK|X) untuk setiap pendaftar beserta status '
         'prediksi akhir dan ringkasan metrik evaluasi model.'),

        ('Interface Rekomendasi.png',
         '6)  Rancangan Halaman Rekomendasi Tindak Lanjut',
         'Gambar 4.23 Rancangan Antarmuka Halaman Rekomendasi Tindak Lanjut',
         'Gambar 4.23 menunjukkan rancangan halaman rekomendasi tindak lanjut yang menampilkan '
         'prioritas follow up bagi setiap calon mahasiswa berdasarkan hasil klasifikasi, sehingga '
         'Tim Marketing dapat menentukan calon mahasiswa mana yang perlu dihubungi terlebih dahulu.'),

        ('Interface Detail Prediksi.png',
         '7)  Rancangan Halaman Detail Prediksi',
         'Gambar 4.24 Rancangan Antarmuka Halaman Detail Prediksi',
         'Gambar 4.24 menunjukkan rancangan halaman detail prediksi yang menampilkan rincian '
         'perhitungan probabilitas (prior, conditional, dan posterior) untuk satu calon mahasiswa '
         'secara lebih rinci beserta hasil prediksi akhirnya.'),

        ('Interface Cetak Export Laporan.png',
         '8)  Rancangan Halaman Cetak/Export Laporan',
         'Gambar 4.25 Rancangan Antarmuka Halaman Cetak/Export Laporan',
         'Gambar 4.25 menunjukkan rancangan halaman laporan yang memungkinkan Tim Marketing '
         'memfilter data berdasarkan rentang tanggal dan status prediksi, kemudian mengekspor '
         'laporan dalam format PDF atau Excel.'),
    ]

    for (img_file, num_heading, caption, explanation) in ui_pages:
        pda_cur = p(num_heading, bold=True, after=pda_cur, ind=1800)
        pda_cur = ins_img_para(pda_cur, img_file, caption, width_cm=15)
        pda_cur = p(explanation, after=pda_cur)

    print("Perancangan Desain Antarmuka inserted.")
else:
    print("WARNING: 'Perancang Design Antar Muka' paragraph not found.")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5b2: PENGKODEAN (screenshot kode program yang penting)
# ═══════════════════════════════════════════════════════════════════════════════
kode_cur = None
for para in doc.paragraphs:
    if para.text.strip() == 'Pengkodean':
        kode_cur = para._element
        break

if kode_cur is not None:
    kode_cur = p('Tahap pengkodean merupakan tahap penerjemahan rancangan sistem dan algoritma '
                 'Naive Bayes ke dalam bahasa pemrograman PHP dengan framework Laravel. Berikut '
                 'adalah beberapa bagian kode program yang menjadi inti dari sistem rekomendasi '
                 'tindak lanjut konfirmasi pendaftaran mahasiswa baru:', after=kode_cur)

    code_snippets = [
        ('Kode - NaiveBayesService train.png', 15,
         '1)  Kelas NaiveBayesService — Fungsi train()',
         'Gambar 4.26 Kode Program Perhitungan Prior dan Conditional Probability',
         'Gambar 4.26 menunjukkan kode program fungsi train() yang menghitung prior '
         'probability dan conditional probability dengan Laplace Smoothing dari data '
         'training.'),

        ('Kode - NaiveBayesService predict.png', 15,
         '2)  Kelas NaiveBayesService — Fungsi predict()',
         'Gambar 4.27 Kode Program Perhitungan Posterior Probability dan Klasifikasi',
         'Gambar 4.27 menunjukkan kode program fungsi predict() yang menghitung posterior '
         'probability P(Y|X) untuk menentukan kelas prediksi (MASUK atau TIDAK MASUK) pada '
         'satu data pendaftar.'),

        ('Kode - NaiveBayesService classifyAll.png', 15,
         '3)  Kelas NaiveBayesService — Fungsi classifyAll()',
         'Gambar 4.28 Kode Program Klasifikasi Massal',
         'Gambar 4.28 menunjukkan kode program fungsi classifyAll() yang melatih model dari '
         'seluruh data historis berlabel kemudian menerapkan prediksi ke seluruh data '
         'pendaftar (historis maupun baru).'),

        ('Kode - NaiveBayesService evaluate 1.png', 15,
         '4)  Kelas NaiveBayesService — Fungsi evaluate() bagian 1',
         'Gambar 4.29 Kode Program Split Data Training dan Testing',
         'Gambar 4.29 menunjukkan bagian awal fungsi evaluate() yang melakukan split data '
         '80:20 secara stratified dengan seed tetap, kemudian melatih model dari data '
         'training.'),

        ('Kode - NaiveBayesService evaluate 2.png', 15,
         '5)  Kelas NaiveBayesService — Fungsi evaluate() bagian 2',
         'Gambar 4.30 Kode Program Confusion Matrix dan Metrik Evaluasi',
         'Gambar 4.30 menunjukkan bagian akhir fungsi evaluate() yang menguji model ke data '
         'testing, menghitung confusion matrix (TP, TN, FP, FN), serta metrik akurasi, '
         'presisi, recall, dan F1-Score.'),

        ('Kode - Migration Pendaftars.png', 15,
         '6)  Migration Tabel Pendaftars',
         'Gambar 4.31 Kode Program Struktur Tabel Basis Data Pendaftar',
         'Gambar 4.31 menunjukkan struktur tabel pendaftars pada basis data MySQL yang '
         'menyimpan seluruh atribut data calon mahasiswa beserta kolom hasil prediksi '
         '(prob_masuk, prob_tidak_masuk, prediksi).'),

        ('Kode - PendaftarController.png', 15,
         '7)  Kelas PendaftarController',
         'Gambar 4.32 Kode Program Pengelolaan Data Pendaftar',
         'Gambar 4.32 menunjukkan kode program controller yang menangani proses tambah dan '
         'ubah data pendaftar. Setiap kali data ditambah atau diubah, fungsi reklasifikasi() '
         'otomatis dipanggil untuk melatih ulang model dan memperbarui hasil prediksi '
         'pendaftar tersebut.'),
    ]

    for (img_file, width_cm, num_heading, caption, explanation) in code_snippets:
        kode_cur = p(num_heading, bold=True, after=kode_cur, ind=1800)
        kode_cur = ins_img_para(kode_cur, img_file, caption, width_cm=width_cm,
                                 source='Sumber: Hasil Implementasi (2026)')
        kode_cur = p(explanation, after=kode_cur)

    print("Pengkodean inserted.")
else:
    print("WARNING: 'Pengkodean' paragraph not found.")

# ═══════════════════════════════════════════════════════════════════════════════
# STEP 5c: PROTOTYPE APLIKASI (screenshot aplikasi Laravel yang sudah dibangun)
# ═══════════════════════════════════════════════════════════════════════════════
proto_cur = None
for para in doc.paragraphs:
    if para.text.strip() == 'Prototype Aplikasi':
        proto_cur = para._element
        break

if proto_cur is not None:
    proto_cur = p('Berdasarkan rancangan desain antarmuka pada sub bab sebelumnya, prototype '
                  'aplikasi kemudian dibangun menggunakan bahasa pemrograman PHP dengan framework '
                  'Laravel dan basis data MySQL. Berikut adalah tampilan hasil implementasi '
                  'prototype sistem rekomendasi tindak lanjut konfirmasi pendaftaran mahasiswa baru:',
                  after=proto_cur)

    app_screenshots = [
        ('App - Login.png',
         '1)  Halaman Login',
         'Gambar 4.33 Tampilan Halaman Login',
         'Gambar 4.33 menunjukkan tampilan halaman login pada prototype yang telah dibangun. '
         'Tim Marketing memasukkan email dan password untuk dapat mengakses sistem.'),

        ('App - Dashboard.png',
         '2)  Halaman Dashboard',
         'Gambar 4.34 Tampilan Halaman Dashboard',
         'Gambar 4.34 menunjukkan tampilan halaman dashboard yang menampilkan ringkasan total '
         'pendaftar, jumlah prediksi MASUK dan TIDAK MASUK, akurasi model, serta daftar prediksi '
         'terbaru.'),

        ('App - Data Pendaftar.png',
         '3)  Halaman Data Pendaftar',
         'Gambar 4.35 Tampilan Halaman Data Pendaftar',
         'Gambar 4.35 menunjukkan tampilan halaman data pendaftar yang menampilkan daftar seluruh '
         'calon mahasiswa beserta fitur pencarian, tambah, ubah, dan hapus data.'),

        ('App - Tambah Data Pendaftar.png',
         '4)  Halaman Tambah Data Pendaftar',
         'Gambar 4.36 Tampilan Halaman Tambah Data Pendaftar',
         'Gambar 4.36 menunjukkan tampilan form input untuk menambah data pendaftar baru sesuai '
         'atribut yang digunakan dalam perhitungan Naive Bayes.'),

        ('App - Hasil Klasifikasi.png',
         '5)  Halaman Hasil Klasifikasi Naive Bayes',
         'Gambar 4.37 Tampilan Halaman Hasil Klasifikasi Naive Bayes',
         'Gambar 4.37 menunjukkan tampilan halaman hasil klasifikasi yang menampilkan nilai '
         'probabilitas P(MASUK|X) dan P(TIDAK MASUK|X) beserta prediksi akhir dan metrik evaluasi '
         'model (akurasi, presisi, recall, F1-Score) yang dihitung secara langsung oleh sistem.'),

        ('App - Detail Prediksi.png',
         '6)  Halaman Detail Prediksi',
         'Gambar 4.38 Tampilan Halaman Detail Prediksi',
         'Gambar 4.38 menunjukkan tampilan halaman detail prediksi yang menampilkan rincian '
         'perhitungan prior probability, conditional probability tiap atribut, dan posterior '
         'probability untuk satu calon mahasiswa secara transparan.'),

        ('App - Rekomendasi.png',
         '7)  Halaman Rekomendasi Tindak Lanjut',
         'Gambar 4.39 Tampilan Halaman Rekomendasi Tindak Lanjut',
         'Gambar 4.39 menunjukkan tampilan halaman rekomendasi tindak lanjut yang menyusun '
         'prioritas follow up bagi setiap calon mahasiswa berdasarkan hasil klasifikasi.'),

        ('App - Laporan.png',
         '8)  Halaman Cetak / Export Laporan',
         'Gambar 4.40 Tampilan Halaman Cetak / Export Laporan',
         'Gambar 4.40 menunjukkan tampilan halaman laporan yang memungkinkan Tim Marketing '
         'memfilter data berdasarkan rentang tanggal dan status prediksi, kemudian mengekspor '
         'laporan dalam format PDF atau Excel.'),
    ]

    for (img_file, num_heading, caption, explanation) in app_screenshots:
        proto_cur = p(num_heading, bold=True, after=proto_cur, ind=1800)
        proto_cur = ins_img_para(proto_cur, img_file, caption, width_cm=15,
                                  source='Sumber: Hasil Implementasi (2026)')
        proto_cur = p(explanation, after=proto_cur)

    print("Prototype Aplikasi inserted.")
else:
    print("WARNING: 'Prototype Aplikasi' paragraph not found.")

# ═══════════════════════════════════════════════════════════════════════════════
# C. PEMBAHASAN — insert after placeholder "Pembahasan" di akhir dokumen
# ═══════════════════════════════════════════════════════════════════════════════
# Cari placeholder "Pembahasan" (bukan "C.  Pembahasan" yang kita buat)
pemb_el = None
for para in doc.paragraphs:
    if para.text.strip() == 'Pembahasan':
        pemb_el = para._element
        break

if pemb_el is not None:
    # Ganti isi placeholder jadi heading "C.  Pembahasan"
    for r in list(pemb_el.findall(qn('w:r'))):
        pemb_el.remove(r)
    pPr = pemb_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); pemb_el.insert(0, pPr)
    sp_el = pPr.find(qn('w:spacing'))
    if sp_el is None:
        sp_el = OxmlElement('w:spacing'); pPr.append(sp_el)
    sp_el.set(qn('w:before'),'0'); sp_el.set(qn('w:after'),'0')
    sp_el.set(qn('w:line'),'360'); sp_el.set(qn('w:lineRule'),'auto')
    r_new = OxmlElement('w:r')
    rPr_new = OxmlElement('w:rPr')
    rf_new = OxmlElement('w:rFonts'); rf_new.set(qn('w:ascii'),TNR); rf_new.set(qn('w:hAnsi'),TNR); rPr_new.append(rf_new)
    rPr_new.append(OxmlElement('w:b'))
    sz_new = OxmlElement('w:sz'); sz_new.set(qn('w:val'),'20'); rPr_new.append(sz_new)
    szc_new = OxmlElement('w:szCs'); szc_new.set(qn('w:val'),'20'); rPr_new.append(szc_new)
    r_new.append(rPr_new)
    t_new = OxmlElement('w:t'); t_new.text = 'C.  Pembahasan'; r_new.append(t_new)
    pemb_el.append(r_new)
    cur = pemb_el
    print("C. Pembahasan placeholder ditemukan dan digunakan.")
else:
    # Fallback: insert setelah paragraf terakhir
    cur = doc.paragraphs[-1]._element
    cur = p('C.  Pembahasan', bold=True, after=cur)
    print("Placeholder tidak ditemukan, insert di akhir dokumen.")

# 1. Pembahasan Hasil Analisis Metode
cur = p('1.  Pembahasan Hasil Analisis Metode', bold=True, after=cur)
cur = p('Penerapan algoritma Naive Bayes dalam penelitian ini bertujuan untuk memberikan '
        'rekomendasi tindak lanjut konfirmasi pendaftaran mahasiswa baru Universitas Tazkia. '
        'Metode ini bekerja dengan memperkirakan kemungkinan status retensi calon mahasiswa '
        'berdasarkan pengalaman dari data historis pendaftaran sebelumnya, yaitu dengan cara '
        'menghitung nilai probabilitas dari setiap atribut pada masing-masing kelas target, '
        'kemudian membandingkan nilai probabilitas akhir untuk menentukan hasil prediksi.', after=cur)

cur = p('Proses klasifikasi dilakukan menggunakan 5 atribut yang telah diseleksi, yaitu '
        'Kategori Jarak Asal (X1), Tingkat Follow Up Internal (X2), Status Test (X3), '
        'Kategori Nilai Test (X4), dan Kategori Penghasilan (X5). Berdasarkan hasil seleksi '
        'fitur pada Tabel 4.3, atribut seperti Kategori Asal Sekolah, Waktu Pendaftaran, dan '
        'Status Uang Pendaftaran tidak diikutsertakan karena tidak memiliki pengaruh signifikan '
        'terhadap penentuan kelas target dalam konteks penelitian ini.', after=cur)

cur = p(f'Berdasarkan hasil perhitungan prior probability pada Tabel 4.5, diketahui bahwa '
        f'proporsi kelas MASUK adalah {str(p_m).replace(".",",")} ({n_m} dari {n_train} data training) dan kelas TIDAK MASUK '
        f'adalah {str(p_t).replace(".",",")} ({n_t} dari {n_train} data training). Komposisi ini menunjukkan bahwa terdapat '
        f'ketidakseimbangan kelas (class imbalance) dimana kelas TIDAK MASUK lebih dominan, '
        f'namun teknik Laplace Smoothing yang diterapkan membantu menjaga konsistensi model.', after=cur)

cur = p(f'Teknik Laplace Smoothing diterapkan pada perhitungan conditional probability untuk '
        f'menghindari nilai probabilitas nol pada nilai atribut yang tidak muncul di data training. '
        f'Penerapan teknik ini terbukti efektif dalam menjaga konsistensi perhitungan posterior '
        f'probability pada seluruh {n_test} data testing, sehingga setiap data testing dapat '
        f'diklasifikasikan dengan hasil yang valid.', after=cur)

# 2. Pembahasan Uji Hasil
cur = p('2.  Pembahasan Uji Hasil', bold=True, after=cur)
cur = p(f'Pada tahap ini dilakukan evaluasi kinerja model menggunakan Confusion Matrix. '
        f'Evaluasi dilakukan dengan membandingkan hasil prediksi algoritma Naive Bayes terhadap '
        f'status aktual pada {n_test} data testing. '
        f'Hasil perbandingan tersebut dapat dilihat pada Tabel 4.13 berikut.', after=cur)

tbl_cm = [
    ['','','Prediksi MASUK','Prediksi TIDAK MASUK'],
    ['Aktual','MASUK',       f'{tp} (TP)', f'{fn} (FN)'],
    ['',      'TIDAK MASUK', f'{fp} (FP)', f'{tn} (TN)'],
]
cur = ins_tbl(cur, tbl_cm, 'Tabel 4.13 Confusion Matrix')

cur = p('Berdasarkan Tabel 4.13, maka berikut adalah hasil evaluasi model yang dihitung '
        'menggunakan rumus di bawah ini.', after=cur)

total = tp + tn + fp + fn

cur = eq_ins([
    mt('Akurasi = '),
    mfrac([mt('TP + TN')], [mt('TP + FP + TN + FN')]),
    mt(' = '),
    mfrac([mt(f'{tp+tn}')], [mt(f'{total}')]),
    mt(f' × 100% = {acc}%')
], after_el=cur)

cur = eq_ins([
    mt('Presisi = '),
    mfrac([mt('TP')], [mt('TP + FP')]),
    mt(' = '),
    mfrac([mt(f'{tp}')], [mt(f'{tp+fp}')]),
    mt(f' × 100% = {prec}%')
], after_el=cur)

cur = eq_ins([
    mt('Recall = '),
    mfrac([mt('TP')], [mt('TP + FN')]),
    mt(' = '),
    mfrac([mt(f'{tp}')], [mt(f'{tp+fn}')]),
    mt(f' × 100% = {rec}%')
], after_el=cur)

cur = eq_ins([
    mt('F1-Score = '),
    mfrac(
        [mt(f'2 × {prec} × {rec}')],
        [mt(f'{prec} + {rec}')]
    ),
    mt(f' = {f1}%')
], after_el=cur)

tbl_ev = [
    ['Metrik','Rumus','Hasil'],
    [{'t':'Akurasi','l':True},  {'t':'(TP+TN) / (TP+FP+TN+FN)','l':True},              f'{acc}%'],
    [{'t':'Presisi','l':True},  {'t':'TP / (TP+FP)','l':True},                          f'{prec}%'],
    [{'t':'Recall','l':True},   {'t':'TP / (TP+FN)','l':True},                          f'{rec}%'],
    [{'t':'F1-Score','l':True}, {'t':'2 x Presisi x Recall / (Presisi+Recall)','l':True}, f'{f1}%'],
]
cur = ins_tbl(cur, tbl_ev, 'Tabel 4.14 Hasil Evaluasi Model Naive Bayes')

benar = tp + tn
salah = fp + fn
cur = p(f'Berdasarkan Tabel 4.13 dan Tabel 4.14, dari {n_test} data testing yang diuji, '
        f'model Naive Bayes berhasil mengklasifikasikan {benar} data dengan benar dan '
        f'{salah} data secara keliru. Rinciannya adalah sebagai berikut: '
        f'{tp} data aktual MASUK diprediksi benar sebagai MASUK (True Positive), '
        f'{tn} data aktual TIDAK MASUK diprediksi benar sebagai TIDAK MASUK (True Negative), '
        f'{fp} data aktual TIDAK MASUK diprediksi keliru sebagai MASUK (False Positive), dan '
        f'{fn} data aktual MASUK diprediksi keliru sebagai TIDAK MASUK (False Negative).', after=cur)

cur = p(f'Nilai recall sebesar {rec}% menunjukkan bahwa model sangat sensitif dalam mendeteksi '
        f'calon mahasiswa yang berpotensi MASUK. Dalam konteks sistem rekomendasi tindak lanjut ini, '
        f'nilai recall yang tinggi sangat diutamakan karena konsekuensi melewatkan calon mahasiswa '
        f'yang berpotensi mendaftar (False Negative) lebih besar dampaknya dibandingkan dengan '
        f'memfollow up calon mahasiswa yang ternyata tidak mendaftar (False Positive).', after=cur)

cur = p(f'Nilai presisi sebesar {prec}% menunjukkan bahwa dari seluruh prediksi MASUK yang dihasilkan '
        f'model, sebagian besar benar-benar berstatus MASUK. Terdapat {fp} kasus False Positive '
        f'yaitu calon mahasiswa yang aktualnya TIDAK MASUK namun diprediksi sebagai MASUK. '
        f'Kasus ini umumnya terjadi pada calon mahasiswa dengan pola atribut yang menyerupai '
        f'kelas MASUK, namun faktor lain di luar atribut penelitian menjadi penentu akhirnya. '
        f'Adapun kasus False Positive ini masih dapat ditoleransi karena dampaknya hanya berupa '
        f'effort follow up tambahan.', after=cur)

cur = p(f'Secara keseluruhan, nilai akurasi model sebesar {acc}% dengan F1-Score {f1}% '
        f'menunjukkan bahwa algoritma Naive Bayes dapat diterapkan dengan baik dalam '
        f'memberikan rekomendasi tindak lanjut konfirmasi pendaftaran mahasiswa baru '
        f'Universitas Tazkia menggunakan dataset sebanyak 1.285 record. '
        f'Hasil ini juga sejalan dengan karakteristik algoritma Naive Bayes yang efektif '
        f'dan efisien untuk klasifikasi dengan atribut-atribut bertipe kategoris '
        f'seperti yang digunakan dalam penelitian ini.', after=cur)

# 3. Analisis Kontribusi Atribut
cur = p('3.  Analisis Kontribusi Atribut', bold=True, after=cur)
cur = p('Berdasarkan hasil perhitungan conditional probability pada Tabel 4.6 hingga Tabel 4.10, '
        'dapat dianalisis kontribusi masing-masing atribut dalam membedakan kelas MASUK dan '
        'TIDAK MASUK sebagai berikut:', after=cur)

pm_sudah = cond_tables['status_test']['MASUK'].get('Sudah Tes', {}).get('prob', 0)
pt_sudah = cond_tables['status_test']['TIDAK MASUK'].get('Sudah Tes', {}).get('prob', 0)
pm_belum = cond_tables['status_test']['MASUK'].get('Belum Tes', {}).get('prob', 0)
pt_belum = cond_tables['status_test']['TIDAK MASUK'].get('Belum Tes', {}).get('prob', 0)
cur = p(f'a)  Atribut Status Test (X3) merupakan atribut yang paling diskriminatif. Calon '
        f'mahasiswa dengan status Sudah Tes memiliki probabilitas MASUK sebesar {str(pm_sudah).replace(".",",")} '
        f'dibandingkan hanya {str(pt_sudah).replace(".",",")} untuk kelas TIDAK MASUK. Sebaliknya, calon mahasiswa '
        f'yang Belum Tes memiliki probabilitas TIDAK MASUK sebesar {str(pt_belum).replace(".",",")} dibandingkan {str(pm_belum).replace(".",",")} '
        f'untuk kelas MASUK. Atribut ini menjadi pembeda terkuat karena mengikuti tes '
        f'merupakan langkah konkret yang mencerminkan keseriusan calon mahasiswa.', after=cur, ind=2520)

pm_nt = cond_tables['kategori_nilai_test']['MASUK'].get('Nilai Tinggi', {}).get('prob', 0)
pt_nt = cond_tables['kategori_nilai_test']['TIDAK MASUK'].get('Nilai Tinggi', {}).get('prob', 0)
cur = p(f'b)  Atribut Kategori Nilai Test (X4) turut memberikan kontribusi signifikan. '
        f'Nilai Tinggi memberikan probabilitas MASUK sebesar {str(pm_nt).replace(".",",")} berbanding {str(pt_nt).replace(".",",")} untuk '
        f'TIDAK MASUK. Namun atribut ini juga berpotensi menyebabkan False Positive, karena '
        f'beberapa calon mahasiswa dengan nilai tinggi ternyata tidak melanjutkan pendaftaran '
        f'karena faktor lain yang tidak tercakup dalam atribut penelitian ini.', after=cur, ind=2520)

cur = p('c)  Atribut Tingkat Follow Up Internal (X2) menunjukkan bahwa sebagian besar calon '
        'mahasiswa yang masuk kategori MASUK berstatus Belum Dihubungi (P=0,6364), '
        'artinya mereka mendaftar secara mandiri tanpa perlu intensitas follow up tinggi. '
        'Sedangkan calon mahasiswa dengan status Follow Up Intensif justru lebih banyak '
        'berakhir sebagai TIDAK MASUK, yang mengindikasikan bahwa intensitas kontak tinggi '
        'belum tentu menjamin konversi pendaftaran.', after=cur, ind=2520)

cur = p('d)  Atribut Kategori Jarak Asal (X1) dan Kategori Penghasilan (X5) memberikan '
        'kontribusi yang lebih kecil dibandingkan atribut lainnya. Meskipun demikian, '
        'kombinasi kedua atribut ini tetap berperan dalam membedakan pola data pada '
        'kasus-kasus tertentu, khususnya ketika nilai atribut lain tidak cukup '
        'diskriminatif untuk menentukan kelas.', after=cur, ind=2520)

# 4. Implikasi Praktis
cur = p('4.  Implikasi Praktis', bold=True, after=cur)
cur = p('Hasil penelitian ini memiliki implikasi praktis yang signifikan bagi Universitas Tazkia, '
        'khususnya bagi tim marketing yang menangani follow up konfirmasi pendaftaran mahasiswa baru. '
        'Dengan menerapkan sistem rekomendasi berbasis Naive Bayes ini, tim marketing dapat '
        'memprioritaskan tindak lanjut kepada calon mahasiswa yang memiliki probabilitas '
        'tinggi untuk MASUK, sehingga alokasi waktu dan sumber daya menjadi lebih efisien '
        'dan tepat sasaran.', after=cur)

cur = p('Secara praktis, sistem ini dapat membantu tim marketing dalam mengidentifikasi '
        'calon mahasiswa berprioritas tinggi berdasarkan kombinasi atribut seperti Status Test '
        'sudah mengikuti ujian, nilai test tinggi, dan jarak domisili yang tidak terlalu jauh. '
        'Calon mahasiswa dengan profil tersebut terbukti memiliki probabilitas MASUK yang '
        'lebih tinggi berdasarkan data historis yang dianalisis dalam penelitian ini.', after=cur)

cur = p('Selain itu, pola yang ditemukan dari kasus False Positive memberikan informasi '
        'berharga bahwa calon mahasiswa yang sudah mengikuti tes dengan nilai tinggi namun '
        'penghasilannya Tidak Diketahui perlu mendapatkan perhatian khusus, misalnya '
        'dengan menggali informasi lebih lanjut terkait faktor penghambat pendaftaran '
        'seperti kendala biaya atau pertimbangan lain yang belum terekam dalam data.', after=cur)

# 5. Keterbatasan dan Saran
cur = p('5.  Keterbatasan dan Saran Pengembangan', bold=True, after=cur)
cur = p(f'Penelitian ini memiliki beberapa keterbatasan yang perlu diperhatikan dalam '
        f'pengembangan ke depan. Pertama, meskipun dataset yang digunakan berjumlah 1.285 record '
        f'dengan pembagian {n_train} data training dan {n_test} data testing, terdapat '
        f'ketidakseimbangan kelas (class imbalance) dimana kelas TIDAK MASUK ({n_t} data) '
        f'jauh lebih banyak dibandingkan kelas MASUK ({n_m} data). Kondisi ini dapat '
        f'memengaruhi kemampuan model dalam memprediksi kelas minoritas secara optimal.', after=cur)

cur = p('Kedua, asumsi independensi antar atribut yang merupakan dasar dari algoritma '
        'Naive Bayes tidak selalu terpenuhi dalam kondisi nyata. Misalnya, atribut '
        'Status Test dan Kategori Nilai Test sangat berkorelasi satu sama lain, '
        'sehingga asumsi ini berpotensi memengaruhi akurasi prediksi pada kasus-kasus tertentu.', after=cur)

cur = p('Untuk pengembangan selanjutnya, disarankan untuk: (1) menggunakan seluruh dataset '
        'yang tersedia dalam proses pelatihan model; (2) membandingkan performa Naive Bayes '
        'dengan algoritma lain seperti Decision Tree atau Random Forest untuk mendapatkan '
        'model terbaik; dan (3) mempertimbangkan penambahan atribut baru yang lebih '
        'representatif seperti riwayat interaksi digital calon mahasiswa atau sumber '
        'informasi pendaftaran, guna meningkatkan akurasi dan keandalan sistem rekomendasi.', after=cur)

doc.save('SUSUN-SKRIPSI-ARIF.docx')
print('SELESAI - File berhasil disimpan.')
