import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn

doc = Document('SUSUN-SKRIPSI-ARIF_BACKUP.docx')

print('=== STYLE paragraf 490-520 (sekitar BAB IV dan isinya) ===')
for i, para in enumerate(doc.paragraphs[488:525], 488):
    style = para.style.name
    txt = para.text.strip()[:60]
    if not txt:
        continue
    font_info = ''
    if para.runs:
        r = para.runs[0]
        font_info = f'font={r.font.name} size={r.font.size} bold={r.bold}'
    pPr = para._element.find(qn('w:pPr'))
    sp_info = ''
    ind_info = ''
    if pPr is not None:
        sp = pPr.find(qn('w:spacing'))
        if sp is not None:
            sp_info = f'line={sp.get(qn("w:line"))} before={sp.get(qn("w:before"))} after={sp.get(qn("w:after"))}'
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            ind_info = f'ind={ind.get(qn("w:left"))}'
    print(f'[{i}] "{style}" | {sp_info} {ind_info} | {txt}')

print()
print('=== CEK PARAGRAF BAB 1 (awal dokumen) ===')
for i, para in enumerate(doc.paragraphs[70:100], 70):
    style = para.style.name
    txt = para.text.strip()[:60]
    if not txt:
        continue
    pPr = para._element.find(qn('w:pPr'))
    sp_info = ''
    if pPr is not None:
        sp = pPr.find(qn('w:spacing'))
        if sp is not None:
            sp_info = f'line={sp.get(qn("w:line"))} before={sp.get(qn("w:before"))} after={sp.get(qn("w:after"))}'
    print(f'[{i}] "{style}" | {sp_info} | {txt}')
